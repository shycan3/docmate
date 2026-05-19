from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT_PATH = Path(__file__).with_name("presentation-overview.docx")


def set_run_font(run, size=10.5, bold=False, color=None):
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    set_run_font(run, bold=bold)


def add_heading(document, text):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    set_run_font(run, size=12, bold=True, color=(34, 88, 154))


def add_body(document, text):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    set_run_font(run, size=9.6)


def build_docx():
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(1.35)
    section.bottom_margin = Cm(1.25)
    section.left_margin = Cm(1.55)
    section.right_margin = Cm(1.55)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("제33회 경영정보시스템 프로젝트 경진대회 발표개요서")
    set_run_font(title_run, size=15, bold=True, color=(20, 54, 96))

    table = document.add_table(rows=3, cols=2)
    table.style = "Table Grid"
    rows = [
        ("팀명", "DocMate"),
        ("작품명", "DocMate: 장학금·청년정책 분석 콘솔"),
        ("작품주제", "공고 문서를 신청 가능성, 위험 조건, 준비 서류, 원문 근거, 실행 체크리스트로 변환하는 로컬 데모 서비스"),
    ]
    for row, (label, value) in zip(table.rows, rows):
        set_cell_text(row.cells[0], label, bold=True)
        set_cell_text(row.cells[1], value)
        shade_cell(row.cells[0], "EAF2F8")

    add_heading(document, "보고서 개요")
    add_body(
        document,
        "DocMate는 장학금과 청년정책 공고를 읽는 데 드는 시간을 줄이고, 사용자가 실제 신청 준비까지 이어갈 수 있도록 돕는 정보시스템 프로젝트이다. 많은 지원 공고는 PDF나 긴 웹 공지 형태로 제공되어 신청 기간, 대상 조건, 소득 요건, 제출 서류, 중복지원 제한을 빠르게 파악하기 어렵다.",
    )
    add_body(
        document,
        "본 프로젝트는 이 문제를 공고 요약이 아니라 신청 준비 의사결정의 문제로 정의했다. 사용자는 나이, 학적, 거주지, 현재 상태, 소득구간, 재학 상태를 입력하고 공고를 선택한다. 시스템은 신청 기간, 지원 대상, 지원 내용, 준비 서류, 신청 방법, URL을 추출한 뒤 사용자 프로필과 비교해 신청 가능 여부를 판정한다.",
    )
    add_body(
        document,
        "구현은 경진대회 발표용 로컬 서비스 완성도를 목표로 했다. Python 기반 API 서버, PyMuPDF PDF 추출, SQLite 히스토리, no-build HTML/CSS/JavaScript 프론트엔드로 구성했으며, 샘플 공고 분석, 원문 근거 표시, 위험 조건 탐지, 체크리스트, Markdown 내보내기, 히스토리 저장/삭제, 공고 비교, 데모 모드를 포함한다.",
    )
    add_body(
        document,
        "개선 과정에서는 사용자·심사위원·미술/웹디자인 계열 페르소나 피드백을 반영했다. 초기 화면을 프로필 작성과 공고 선택 중심으로 단순화했고, 거주지와 직업 같은 항목은 자유 입력 대신 선택형 컨트롤로 바꾸었다. 또한 자동 판정의 신뢰성을 높이기 위해 결과마다 원문 근거를 펼쳐 볼 수 있게 했다.",
    )
    add_body(
        document,
        "발표의 핵심은 세 가지이다. 첫째, 긴 공고를 읽기 전에 신청 가능성과 위험 조건을 먼저 확인하게 만든다. 둘째, 자동 판정의 근거를 원문 스니펫으로 확인하게 한다. 셋째, 분석 결과를 준비 서류와 마감 확인 체크리스트로 바꿔 사용자의 다음 행동을 명확하게 만든다.",
    )

    add_heading(document, "기대효과 및 확장")
    add_body(
        document,
        "기대효과는 정보 탐색 시간 절감, 신청 누락 감소, 정책 접근성 향상, 반복 상담 업무 감소이다. 향후에는 실공고 수집, OCR, 사용자 계정, 개인정보 보호 정책, 캘린더 연동, 기관용 대시보드로 확장할 수 있다.",
    )

    role_table = document.add_table(rows=1, cols=2)
    role_table.style = "Table Grid"
    set_cell_text(role_table.rows[0].cells[0], "역할", bold=True)
    set_cell_text(role_table.rows[0].cells[1], "발표에서 설명할 책임", bold=True)
    shade_cell(role_table.rows[0].cells[0], "EAF2F8")
    shade_cell(role_table.rows[0].cells[1], "EAF2F8")
    for role, desc in [
        ("기획/문제정의", "청년 지원 공고 탐색 문제 정의, 사용자 흐름 설계"),
        ("데이터/분석", "공고 추출 항목 정의, 규칙 기반 판정, 위험 조건 구조화"),
        ("백엔드", "로컬 API, PDF/TXT 분석, SQLite 히스토리, 샘플 API 구현"),
        ("프론트엔드/UX", "프로필 입력, 결과 요약, 원문 근거, 히스토리 비교, 데모 모드 구현"),
        ("발표/검증", "페르소나 피드백 정리, 심사 Q&A, 발표 흐름 검증"),
    ]:
        row = role_table.add_row()
        set_cell_text(row.cells[0], role, bold=True)
        set_cell_text(row.cells[1], desc)

    document.save(OUT_PATH)


if __name__ == "__main__":
    build_docx()
    print(OUT_PATH)
